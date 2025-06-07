import fitz
from docx import Document as DocxDocument
from langchain.schema import Document
from pathlib import Path
from src.logger import get_logger
from src.custom_exception import CustomException

logger = get_logger(__name__)

def load_and_split_document(file_path: str) -> list:
    try:
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix == ".pdf":
            return _load_pdf_with_metadata(file_path)
        elif suffix == ".docx":
            return _load_docx_with_metadata(file_path)
        else:
            raise CustomException("Unsupported file type", f"File: {file_path}")

    except Exception as e:
        logger.error(f"Error while loading and splitting document: {e}")
        raise CustomException("Error while loading and splitting document", e)


def _load_pdf_with_metadata(file_path: str) -> list:
    try:
        logger.info(f"Loading PDF using PyMuPDF: {file_path}")
        doc = fitz.open(file_path)
        documents = []

        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            if text.strip():
                documents.append(Document(
                    page_content=text,
                    metadata={"page": page_num + 1}
                ))
        doc.close()
        logger.info("PDF loaded successfully with page metadata.")
        return documents
    except Exception as e:
        raise CustomException("Error while processing PDF", e)


def _load_docx_with_metadata(file_path: str) -> list:
    try:
        logger.info(f"Loading DOCX using python-docx: {file_path}")
        doc = DocxDocument(file_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        documents = []

        for i, para in enumerate(paragraphs):
            documents.append(Document(
                page_content=para,
                metadata={"section": i + 1}
            ))
        logger.info("DOCX loaded successfully with section metadata.")
        return documents
    except Exception as e:
        raise CustomException("Error while processing DOCX", e)
